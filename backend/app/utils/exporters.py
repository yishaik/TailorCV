"""
Export utilities for generating Markdown, DOCX, and PDF outputs.
"""
import io
from typing import Optional
from ..models.output import TailoredCV, CoverLetter


def generate_markdown(cv: TailoredCV, cover_letter: Optional[CoverLetter] = None) -> str:
    """
    Generate Markdown format CV.
    
    Args:
        cv: Tailored CV data
        cover_letter: Optional cover letter
    
    Returns:
        Markdown formatted string
    """
    lines = []
    
    # Header
    lines.append(f"# {cv.header.name}")
    lines.append(f"**{cv.header.title}**")
    lines.append("")
    
    # Contact info
    contact_parts = []
    if cv.header.contact.get("email"):
        contact_parts.append(cv.header.contact["email"])
    if cv.header.contact.get("phone"):
        contact_parts.append(cv.header.contact["phone"])
    if cv.header.contact.get("location"):
        contact_parts.append(cv.header.contact["location"])
    if cv.header.contact.get("linkedin"):
        contact_parts.append(f"[LinkedIn]({cv.header.contact['linkedin']})")
    
    if contact_parts:
        lines.append(" | ".join(contact_parts))
        lines.append("")
    
    # Summary
    lines.append("## Summary")
    lines.append(cv.summary)
    lines.append("")
    
    # Experience
    if cv.experience:
        lines.append("## Experience")
        lines.append("")
        for exp in cv.experience:
            location_str = f" | {exp.location}" if exp.location else ""
            lines.append(f"### {exp.title}")
            lines.append(f"**{exp.company}** | {exp.dates}{location_str}")
            lines.append("")
            for bullet in exp.bullets:
                lines.append(f"- {bullet.text}")
            lines.append("")
    
    # Skills
    if cv.skills.primary or cv.skills.secondary or cv.skills.tools:
        lines.append("## Skills")
        lines.append("")
        if cv.skills.primary:
            lines.append(f"**Core:** {', '.join(cv.skills.primary)}")
        if cv.skills.secondary:
            lines.append(f"**Additional:** {', '.join(cv.skills.secondary)}")
        if cv.skills.tools:
            lines.append(f"**Tools & Technologies:** {', '.join(cv.skills.tools)}")
        lines.append("")
    
    # Education
    if cv.education:
        lines.append("## Education")
        lines.append("")
        for edu in cv.education:
            year_str = f" ({edu.year})" if edu.year else ""
            lines.append(f"**{edu.degree} in {edu.field}**{year_str}")
            lines.append(f"{edu.institution}")
            for highlight in edu.highlights:
                lines.append(f"- {highlight}")
            lines.append("")
    
    # Certifications
    if cv.certifications:
        lines.append("## Certifications")
        lines.append("")
        for cert in cv.certifications:
            date_str = f" ({cert.date})" if cert.date else ""
            lines.append(f"- **{cert.name}** - {cert.issuer}{date_str}")
        lines.append("")
    
    # Projects
    if cv.projects:
        lines.append("## Projects")
        lines.append("")
        for proj in cv.projects:
            lines.append(f"### {proj.name}")
            lines.append(proj.description)
            if proj.technologies:
                lines.append(f"*Technologies: {', '.join(proj.technologies)}*")
            lines.append("")
    
    cv_markdown = "\n".join(lines)
    
    # Add cover letter if present
    if cover_letter:
        cl_lines = [
            "",
            "---",
            "",
            "# Cover Letter",
            "",
            cover_letter.full_text
        ]
        cv_markdown += "\n".join(cl_lines)
    
    return cv_markdown


def generate_docx(cv: TailoredCV, cover_letter: Optional[CoverLetter] = None) -> bytes:
    """
    Generate DOCX format CV.
    
    Args:
        cv: Tailored CV data
        cover_letter: Optional cover letter
    
    Returns:
        DOCX file content as bytes
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Set narrow margins
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    # Header - Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(cv.header.name)
    name_run.bold = True
    name_run.font.size = Pt(18)
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(cv.header.title)
    title_run.font.size = Pt(12)
    
    # Contact info
    contact_parts = []
    if cv.header.contact.get("email"):
        contact_parts.append(cv.header.contact["email"])
    if cv.header.contact.get("phone"):
        contact_parts.append(cv.header.contact["phone"])
    if cv.header.contact.get("location"):
        contact_parts.append(cv.header.contact["location"])
    
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.add_run(" | ".join(contact_parts))
    
    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(cv.summary)
    
    # Experience
    if cv.experience:
        doc.add_heading("Experience", level=1)
        for exp in cv.experience:
            # Company and title
            exp_para = doc.add_paragraph()
            title_run = exp_para.add_run(f"{exp.title} at {exp.company}")
            title_run.bold = True
            
            # Date and location
            date_para = doc.add_paragraph()
            date_text = exp.dates
            if exp.location:
                date_text += f" | {exp.location}"
            date_para.add_run(date_text).italic = True
            
            # Bullets
            for bullet in exp.bullets:
                bullet_para = doc.add_paragraph(bullet.text, style='List Bullet')
    
    # Skills
    if cv.skills.primary or cv.skills.secondary or cv.skills.tools:
        doc.add_heading("Skills", level=1)
        if cv.skills.primary:
            para = doc.add_paragraph()
            para.add_run("Core: ").bold = True
            para.add_run(", ".join(cv.skills.primary))
        if cv.skills.secondary:
            para = doc.add_paragraph()
            para.add_run("Additional: ").bold = True
            para.add_run(", ".join(cv.skills.secondary))
        if cv.skills.tools:
            para = doc.add_paragraph()
            para.add_run("Tools & Technologies: ").bold = True
            para.add_run(", ".join(cv.skills.tools))
    
    # Education
    if cv.education:
        doc.add_heading("Education", level=1)
        for edu in cv.education:
            edu_para = doc.add_paragraph()
            degree_text = f"{edu.degree} in {edu.field}"
            if edu.year:
                degree_text += f" ({edu.year})"
            edu_para.add_run(degree_text).bold = True
            doc.add_paragraph(edu.institution)
            for highlight in edu.highlights:
                doc.add_paragraph(highlight, style='List Bullet')
    
    # Certifications
    if cv.certifications:
        doc.add_heading("Certifications", level=1)
        for cert in cv.certifications:
            cert_text = f"{cert.name} - {cert.issuer}"
            if cert.date:
                cert_text += f" ({cert.date})"
            doc.add_paragraph(cert_text, style='List Bullet')
    
    # Projects
    if cv.projects:
        doc.add_heading("Projects", level=1)
        for proj in cv.projects:
            proj_para = doc.add_paragraph()
            proj_para.add_run(proj.name).bold = True
            doc.add_paragraph(proj.description)
            if proj.technologies:
                tech_para = doc.add_paragraph()
                tech_para.add_run("Technologies: ").italic = True
                tech_para.add_run(", ".join(proj.technologies))
    
    # Cover letter on new page
    if cover_letter:
        doc.add_page_break()
        doc.add_heading("Cover Letter", level=1)
        doc.add_paragraph(cover_letter.hook)
        doc.add_paragraph(cover_letter.value_proposition)
        doc.add_paragraph(cover_letter.fit_narrative)
        doc.add_paragraph(cover_letter.closing)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_pdf(cv: TailoredCV, cover_letter: Optional[CoverLetter] = None) -> bytes:
    """
    Generate PDF format CV using a bundled TrueType font (Vercel-safe).

    Args:
        cv: Tailored CV data
        cover_letter: Optional cover letter

    Returns:
        PDF file content as bytes
    """
    from fpdf import FPDF
    from pathlib import Path

    fonts_dir = Path(__file__).resolve().parents[1] / "assets" / "fonts"
    regular_font = fonts_dir / "DejaVuSans.ttf"
    bold_font = fonts_dir / "DejaVuSans-Bold.ttf"
    if not regular_font.exists() or not bold_font.exists():
        raise FileNotFoundError("PDF fonts are missing from backend/app/assets/fonts")

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.add_font("DejaVu", "", str(regular_font), uni=True)
    pdf.add_font("DejaVu", "B", str(bold_font), uni=True)

    def add_centered(text: str, size: int, style: str = "") -> None:
        pdf.set_font("DejaVu", style, size)
        pdf.cell(0, 8, text, ln=True, align="C")

    def add_section(title: str) -> None:
        pdf.ln(2)
        pdf.set_font("DejaVu", "B", 12)
        pdf.cell(0, 7, title, ln=True)

    def add_paragraph(text: str) -> None:
        pdf.set_font("DejaVu", "", 10)
        pdf.multi_cell(0, 5, text)

    def add_bullets(items: list[str]) -> None:
        pdf.set_font("DejaVu", "", 10)
        for item in items:
            pdf.multi_cell(0, 5, f"- {item}")

    # Header
    add_centered(cv.header.name, 18, "B")
    add_centered(cv.header.title, 12)

    contact_parts = []
    if cv.header.contact.get("email"):
        contact_parts.append(cv.header.contact["email"])
    if cv.header.contact.get("phone"):
        contact_parts.append(cv.header.contact["phone"])
    if cv.header.contact.get("location"):
        contact_parts.append(cv.header.contact["location"])
    if contact_parts:
        pdf.set_font("DejaVu", "", 9)
        pdf.multi_cell(0, 5, " | ".join(contact_parts), align="C")

    # Summary
    add_section("Summary")
    add_paragraph(cv.summary)

    # Experience
    if cv.experience:
        add_section("Experience")
        for exp in cv.experience:
            pdf.set_font("DejaVu", "B", 11)
            pdf.cell(0, 6, exp.title, ln=True)
            pdf.set_font("DejaVu", "", 10)
            location_str = f" | {exp.location}" if exp.location else ""
            pdf.cell(0, 5, f"{exp.company} | {exp.dates}{location_str}", ln=True)
            add_bullets([bullet.text for bullet in exp.bullets])
            pdf.ln(1)

    # Skills
    if cv.skills.primary or cv.skills.secondary or cv.skills.tools:
        add_section("Skills")
        if cv.skills.primary:
            add_paragraph(f"Core: {', '.join(cv.skills.primary)}")
        if cv.skills.secondary:
            add_paragraph(f"Additional: {', '.join(cv.skills.secondary)}")
        if cv.skills.tools:
            add_paragraph(f"Tools & Technologies: {', '.join(cv.skills.tools)}")

    # Education
    if cv.education:
        add_section("Education")
        for edu in cv.education:
            pdf.set_font("DejaVu", "B", 10)
            year_str = f" ({edu.year})" if edu.year else ""
            pdf.cell(0, 5, f"{edu.degree} in {edu.field}{year_str}", ln=True)
            add_paragraph(edu.institution)
            add_bullets(edu.highlights)
            pdf.ln(1)

    # Certifications
    if cv.certifications:
        add_section("Certifications")
        for cert in cv.certifications:
            date_str = f" ({cert.date})" if cert.date else ""
            add_paragraph(f"{cert.name} - {cert.issuer}{date_str}")

    # Projects
    if cv.projects:
        add_section("Projects")
        for proj in cv.projects:
            pdf.set_font("DejaVu", "B", 10)
            pdf.cell(0, 5, proj.name, ln=True)
            add_paragraph(proj.description)
            if proj.technologies:
                add_paragraph(f"Technologies: {', '.join(proj.technologies)}")
            pdf.ln(1)

    # Cover letter on new page
    if cover_letter:
        pdf.add_page()
        add_section("Cover Letter")
        add_paragraph(cover_letter.hook)
        add_paragraph(cover_letter.value_proposition)
        add_paragraph(cover_letter.fit_narrative)
        add_paragraph(cover_letter.closing)

    pdf_bytes = pdf.output(dest="S")
    if isinstance(pdf_bytes, str):
        pdf_bytes = pdf_bytes.encode("latin-1")
    return pdf_bytes
