"""
Document parsing utilities for PDF, DOCX, and plain text.
"""
import io
from typing import Optional
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from PDF content.
    
    Args:
        content: PDF file content as bytes
    
    Returns:
        Extracted text
    """
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    
    except Exception as e:
        logger.error(f"Failed to extract PDF text: {e}")
        raise ValueError(f"Could not parse PDF: {e}")


def extract_text_from_docx(content: bytes) -> str:
    """
    Extract text from DOCX content.
    
    Args:
        content: DOCX file content as bytes
    
    Returns:
        Extracted text
    """
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    except Exception as e:
        logger.error(f"Failed to extract DOCX text: {e}")
        raise ValueError(f"Could not parse DOCX: {e}")


def extract_text(content: bytes, filename: str) -> str:
    """
    Extract text from a document based on file extension.
    
    Args:
        content: File content as bytes
        filename: Original filename with extension
    
    Returns:
        Extracted text
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(content)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(content)
    elif filename_lower.endswith(".doc"):
        raise ValueError("Legacy .doc format not supported. Please convert to .docx")
    elif filename_lower.endswith((".txt", ".md")):
        return content.decode("utf-8", errors="replace")
    else:
        # Try to decode as plain text
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:
            raise ValueError(f"Unsupported file format: {filename}")


def clean_extracted_text(text: str) -> str:
    """
    Clean up extracted text for better parsing.
    
    Args:
        text: Raw extracted text
    
    Returns:
        Cleaned text
    """
    import re
    
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Fix common OCR/extraction issues
    text = text.replace('•', '\n• ')
    text = text.replace('●', '\n• ')
    text = text.replace('○', '\n• ')
    
    # Restore paragraph breaks at common CV section headers
    section_headers = [
        'experience', 'education', 'skills', 'summary', 'objective',
        'certifications', 'projects', 'publications', 'languages',
        'references', 'awards', 'achievements', 'volunteer'
    ]
    
    for header in section_headers:
        # Case-insensitive replacement
        pattern = re.compile(rf'\s({header})\s', re.IGNORECASE)
        text = pattern.sub(rf'\n\n\1\n', text)
    
    # Clean up extra newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()
